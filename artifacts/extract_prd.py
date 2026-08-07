from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\340710\Downloads\通用能碳平台产品PRD_整合版_一期研发评审稿_V1.0.docx")
OUTPUT = Path(r"D:\Project\Generic_Energy_Carbon_Platform\artifacts\prd-extracted.md")


def iter_blocks(parent):
    if isinstance(parent, DocumentType):
        element = parent.element.body
    else:
        element = parent._tc
    for child in element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


document = Document(SOURCE)
lines = [
    f"# {SOURCE.stem}",
    "",
    f"- 段落数：{len(document.paragraphs)}",
    f"- 表格数：{len(document.tables)}",
    "",
]

for block in iter_blocks(document):
    if isinstance(block, Paragraph):
        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style else ""
        if style.startswith("Heading"):
            try:
                level = min(6, max(1, int(style.split()[-1])))
            except ValueError:
                level = 2
            lines.extend([f"{'#' * level} {text}", ""])
        else:
            lines.extend([text, ""])
    else:
        rows = [[cell.text.replace("\n", " / ").strip() for cell in row.cells] for row in block.rows]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")

OUTPUT.write_text("\n".join(lines), encoding="utf-8")
print(f"{OUTPUT}\nparagraphs={len(document.paragraphs)} tables={len(document.tables)}")
